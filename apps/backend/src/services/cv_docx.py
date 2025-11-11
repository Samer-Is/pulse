"""CV generation service using python-docx."""

from docx import Document
from docx.shared import Pt, RGBColor
from typing import List
import io

from ..schemas.cv import CVGenerateRequest, CVEducation, CVExperience
from ..core.logging import get_logger

logger = get_logger(__name__)


def generate_cv_docx(request: CVGenerateRequest) -> io.BytesIO:
    """
    Generate CV as DOCX file.
    
    Args:
        request: CV generation request with all data
    
    Returns:
        BytesIO buffer containing DOCX file
    """
    doc = Document()
    
    # Header - Name and Title
    heading = doc.add_heading(request.full_name, 0)
    heading.alignment = 1 if request.locale == "ar" else 0  # RTL for Arabic
    
    title = doc.add_paragraph(request.title)
    title.alignment = 1 if request.locale == "ar" else 0
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.color.rgb = RGBColor(0, 0, 128)
    
    # Contact Information
    contact = doc.add_paragraph()
    contact.alignment = 1 if request.locale == "ar" else 0
    contact.add_run(f"📧 {request.email}  ")
    if request.phone:
        contact.add_run(f"📱 {request.phone}")
    
    doc.add_paragraph()  # Spacing
    
    # Summary
    doc.add_heading("Professional Summary" if request.locale == "en" else "الملخص المهني", level=1)
    doc.add_paragraph(request.summary)
    
    # Skills
    if request.skills:
        doc.add_heading("Skills" if request.locale == "en" else "المهارات", level=1)
        skills_para = doc.add_paragraph()
        skills_para.add_run(", ".join(request.skills))
    
    # Experience
    if request.experience:
        doc.add_heading("Experience" if request.locale == "en" else "الخبرة العملية", level=1)
        for exp in request.experience:
            exp_title = doc.add_paragraph(style='List Bullet')
            exp_title.add_run(f"{exp.title} - {exp.company}").bold = True
            
            exp_details = doc.add_paragraph()
            exp_details.add_run(f"{exp.location} | {exp.start_date} - {exp.end_date or 'Present'}")
            exp_details.runs[0].italic = True
            
            if exp.description:
                doc.add_paragraph(exp.description)
    
    # Education
    if request.education:
        doc.add_heading("Education" if request.locale == "en" else "التعليم", level=1)
        for edu in request.education:
            edu_title = doc.add_paragraph(style='List Bullet')
            edu_title.add_run(f"{edu.degree} - {edu.institution}").bold = True
            
            edu_details = doc.add_paragraph()
            edu_details.add_run(f"{edu.location} | {edu.start_date} - {edu.end_date or 'Present'}")
            edu_details.runs[0].italic = True
            
            if edu.description:
                doc.add_paragraph(edu.description)
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    logger.info(f"CV generated: name={request.full_name}, locale={request.locale}")
    
    return buffer

