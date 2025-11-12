"""CV Generator Service - Creates DOCX and PDF resumes."""

import io
import os
from typing import Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from playwright.async_api import async_playwright

from ..schemas.cv import CVRequest, PersonalInfo, Experience, Education, Skill


class CVGenerator:
    """Generate CVs in DOCX and PDF formats."""

    def __init__(self):
        """Initialize CV generator."""
        self.default_font = "Calibri"
        self.heading_color = RGBColor(31, 78, 121)  # Professional blue

    async def generate_docx(self, cv_data: CVRequest) -> bytes:
        """
        Generate CV in DOCX format.
        
        Args:
            cv_data: CV data
            
        Returns:
            DOCX file as bytes
        """
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Personal Information Header
        self._add_personal_info(doc, cv_data.personal_info)
        
        # Summary
        if cv_data.summary:
            self._add_summary(doc, cv_data.summary)
        
        # Experience
        if cv_data.experience:
            self._add_experience(doc, cv_data.experience)
        
        # Education
        if cv_data.education:
            self._add_education(doc, cv_data.education)
        
        # Skills
        if cv_data.skills:
            self._add_skills(doc, cv_data.skills)
        
        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    def _add_personal_info(self, doc: Document, info: PersonalInfo) -> None:
        """Add personal information header."""
        # Name (Large, Bold)
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(info.full_name)
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.color.rgb = self.heading_color
        
        # Contact Info (Centered)
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = [info.email]
        if info.phone:
            contact_parts.append(info.phone)
        if info.location:
            contact_parts.append(info.location)
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(10)
        
        # Links (Centered)
        if info.website or info.linkedin or info.github:
            links_para = doc.add_paragraph()
            links_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            link_parts = []
            if info.website:
                link_parts.append(info.website)
            if info.linkedin:
                link_parts.append(f"LinkedIn: {info.linkedin}")
            if info.github:
                link_parts.append(f"GitHub: {info.github}")
            links_run = links_para.add_run(" | ".join(link_parts))
            links_run.font.size = Pt(9)
            links_run.font.color.rgb = RGBColor(100, 100, 100)
        
        # Add spacing
        doc.add_paragraph()

    def _add_summary(self, doc: Document, summary: str) -> None:
        """Add professional summary section."""
        # Section heading
        heading = doc.add_heading("Professional Summary", level=1)
        heading.runs[0].font.color.rgb = self.heading_color
        
        # Summary text
        summary_para = doc.add_paragraph(summary)
        summary_para.paragraph_format.space_after = Pt(12)

    def _add_experience(self, doc: Document, experiences: list[Experience]) -> None:
        """Add work experience section."""
        # Section heading
        heading = doc.add_heading("Experience", level=1)
        heading.runs[0].font.color.rgb = self.heading_color
        
        for exp in experiences:
            # Job title and company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"{exp.job_title}")
            title_run.font.bold = True
            title_run.font.size = Pt(12)
            
            # Company and dates
            details_para = doc.add_paragraph()
            company_run = details_para.add_run(f"{exp.company}")
            company_run.font.italic = True
            
            if exp.location:
                details_para.add_run(f", {exp.location}")
            
            date_str = exp.start_date
            if exp.end_date:
                date_str += f" - {exp.end_date}"
            details_para.add_run(f" | {date_str}")
            details_para.runs[-1].font.color.rgb = RGBColor(100, 100, 100)
            
            # Description
            if exp.description:
                desc_para = doc.add_paragraph(exp.description)
                desc_para.paragraph_format.space_after = Pt(6)
            
            # Responsibilities (bullet points)
            if exp.responsibilities:
                for resp in exp.responsibilities:
                    bullet_para = doc.add_paragraph(resp, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
            
            # Add spacing between jobs
            doc.add_paragraph()

    def _add_education(self, doc: Document, education: list[Education]) -> None:
        """Add education section."""
        # Section heading
        heading = doc.add_heading("Education", level=1)
        heading.runs[0].font.color.rgb = self.heading_color
        
        for edu in education:
            # Degree
            degree_para = doc.add_paragraph()
            degree_run = degree_para.add_run(edu.degree)
            degree_run.font.bold = True
            degree_run.font.size = Pt(12)
            
            # Institution and dates
            details_para = doc.add_paragraph()
            inst_run = details_para.add_run(edu.institution)
            inst_run.font.italic = True
            
            if edu.location:
                details_para.add_run(f", {edu.location}")
            
            if edu.start_date or edu.end_date:
                date_str = edu.start_date or ""
                if edu.end_date:
                    date_str += f" - {edu.end_date}" if date_str else edu.end_date
                details_para.add_run(f" | {date_str}")
                details_para.runs[-1].font.color.rgb = RGBColor(100, 100, 100)
            
            # GPA
            if edu.gpa:
                gpa_para = doc.add_paragraph(f"GPA: {edu.gpa}")
            
            # Achievements
            if edu.achievements:
                for achievement in edu.achievements:
                    bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
            
            # Add spacing
            doc.add_paragraph()

    def _add_skills(self, doc: Document, skills: list[Skill]) -> None:
        """Add skills section."""
        # Section heading
        heading = doc.add_heading("Skills", level=1)
        heading.runs[0].font.color.rgb = self.heading_color
        
        for skill_group in skills:
            skill_para = doc.add_paragraph()
            
            # Category (if provided)
            if skill_group.category:
                category_run = skill_para.add_run(f"{skill_group.category}: ")
                category_run.font.bold = True
            
            # Skills list
            skills_text = ", ".join(skill_group.skills)
            skill_para.add_run(skills_text)

    async def generate_pdf(self, cv_data: CVRequest) -> bytes:
        """
        Generate CV in PDF format using Playwright.
        
        Args:
            cv_data: CV data
            
        Returns:
            PDF file as bytes
        """
        # Generate HTML
        html_content = self._generate_html(cv_data)
        
        # Convert HTML to PDF using Playwright
        async with async_playwright() as p:
            browser = await p.chromium.launch()
            page = await browser.new_page()
            await page.set_content(html_content)
            pdf_bytes = await page.pdf(
                format='A4',
                margin={
                    'top': '0.5in',
                    'right': '0.75in',
                    'bottom': '0.5in',
                    'left': '0.75in',
                }
            )
            await browser.close()
        
        return pdf_bytes

    def _generate_html(self, cv_data: CVRequest) -> str:
        """Generate HTML for PDF conversion."""
        html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        body {{
            font-family: 'Calibri', 'Arial', sans-serif;
            font-size: 11pt;
            line-height: 1.6;
            color: #333;
        }}
        h1 {{
            color: #1f4e79;
            font-size: 24pt;
            text-align: center;
            margin-bottom: 5px;
        }}
        h2 {{
            color: #1f4e79;
            font-size: 14pt;
            border-bottom: 2px solid #1f4e79;
            padding-bottom: 3px;
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        .contact {{
            text-align: center;
            font-size: 10pt;
            margin-bottom: 5px;
        }}
        .links {{
            text-align: center;
            font-size: 9pt;
            color: #666;
            margin-bottom: 20px;
        }}
        .job-title {{
            font-weight: bold;
            font-size: 12pt;
        }}
        .company {{
            font-style: italic;
        }}
        .date {{
            color: #666;
        }}
        ul {{
            margin: 5px 0;
            padding-left: 20px;
        }}
        li {{
            margin-bottom: 3px;
        }}
        .section-item {{
            margin-bottom: 15px;
        }}
    </style>
</head>
<body>
"""
        
        # Personal Info
        html += f"<h1>{cv_data.personal_info.full_name}</h1>\n"
        
        contact_parts = [cv_data.personal_info.email]
        if cv_data.personal_info.phone:
            contact_parts.append(cv_data.personal_info.phone)
        if cv_data.personal_info.location:
            contact_parts.append(cv_data.personal_info.location)
        html += f"<div class='contact'>{' | '.join(contact_parts)}</div>\n"
        
        if cv_data.personal_info.website or cv_data.personal_info.linkedin or cv_data.personal_info.github:
            link_parts = []
            if cv_data.personal_info.website:
                link_parts.append(cv_data.personal_info.website)
            if cv_data.personal_info.linkedin:
                link_parts.append(f"LinkedIn: {cv_data.personal_info.linkedin}")
            if cv_data.personal_info.github:
                link_parts.append(f"GitHub: {cv_data.personal_info.github}")
            html += f"<div class='links'>{' | '.join(link_parts)}</div>\n"
        
        # Summary
        if cv_data.summary:
            html += "<h2>Professional Summary</h2>\n"
            html += f"<p>{cv_data.summary}</p>\n"
        
        # Experience
        if cv_data.experience:
            html += "<h2>Experience</h2>\n"
            for exp in cv_data.experience:
                html += "<div class='section-item'>\n"
                html += f"<div class='job-title'>{exp.job_title}</div>\n"
                
                company_line = f"<span class='company'>{exp.company}</span>"
                if exp.location:
                    company_line += f", {exp.location}"
                
                date_str = exp.start_date
                if exp.end_date:
                    date_str += f" - {exp.end_date}"
                company_line += f" | <span class='date'>{date_str}</span>"
                
                html += f"<div>{company_line}</div>\n"
                
                if exp.description:
                    html += f"<p>{exp.description}</p>\n"
                
                if exp.responsibilities:
                    html += "<ul>\n"
                    for resp in exp.responsibilities:
                        html += f"<li>{resp}</li>\n"
                    html += "</ul>\n"
                
                html += "</div>\n"
        
        # Education
        if cv_data.education:
            html += "<h2>Education</h2>\n"
            for edu in cv_data.education:
                html += "<div class='section-item'>\n"
                html += f"<div class='job-title'>{edu.degree}</div>\n"
                
                inst_line = f"<span class='company'>{edu.institution}</span>"
                if edu.location:
                    inst_line += f", {edu.location}"
                
                if edu.start_date or edu.end_date:
                    date_str = edu.start_date or ""
                    if edu.end_date:
                        date_str += f" - {edu.end_date}" if date_str else edu.end_date
                    inst_line += f" | <span class='date'>{date_str}</span>"
                
                html += f"<div>{inst_line}</div>\n"
                
                if edu.gpa:
                    html += f"<p>GPA: {edu.gpa}</p>\n"
                
                if edu.achievements:
                    html += "<ul>\n"
                    for achievement in edu.achievements:
                        html += f"<li>{achievement}</li>\n"
                    html += "</ul>\n"
                
                html += "</div>\n"
        
        # Skills
        if cv_data.skills:
            html += "<h2>Skills</h2>\n"
            for skill_group in cv_data.skills:
                if skill_group.category:
                    html += f"<p><strong>{skill_group.category}:</strong> {', '.join(skill_group.skills)}</p>\n"
                else:
                    html += f"<p>{', '.join(skill_group.skills)}</p>\n"
        
        html += """
</body>
</html>
"""
        return html

